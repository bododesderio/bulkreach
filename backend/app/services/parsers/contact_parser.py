# @author Bodo Desderio <rooiboktechltd@gmail.com>
# @copyright 2026 Rooibok Technologies. All rights reserved.
"""Contact import parser (Sections 3.1, 6.2).

Accepts CSV / Excel / Word / PDF / pasted text and returns a ParseResult:
normalised contacts (phone E.164, email, raw_data), detected phone/email columns,
merge-tag column names, and valid / duplicate / error counts (up to 50 messages).
"""
from __future__ import annotations

import csv
import io
import re
from dataclasses import dataclass, field
from typing import Any

import chardet
import pandas as pd

# --- Column detection hints (Section 3.1.1) ---
_PHONE_HINTS = {"phone", "mobile", "tel", "telephone", "msisdn", "cell", "contact", "number", "phone_number"}
_EMAIL_HINTS = {"email", "e-mail", "email_address", "mail", "e_mail"}

_EMAIL_RE = re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")
_MAX_ERRORS = 50
# Hard ceiling on rows materialised from a single upload (zip-bomb / OOM guard).
_MAX_ROWS = 100_000


@dataclass
class ParseResult:
    contacts: list[dict[str, Any]] = field(default_factory=list)  # {phone, email, raw_data}
    columns: list[str] = field(default_factory=list)
    phone_column: str | None = None
    email_column: str | None = None
    merge_columns: list[str] = field(default_factory=list)
    total_rows: int = 0
    valid_contacts: int = 0
    duplicate_count: int = 0
    error_count: int = 0
    errors: list[str] = field(default_factory=list)


# --- Normalisation ---------------------------------------------------------
def normalise_phone(raw: Any, default_cc: str = "256") -> str | None:
    """Normalise a Ugandan/international number to E.164, or None if invalid."""
    if raw is None:
        return None
    digits = re.sub(r"[^\d+]", "", str(raw).strip())
    if not digits:
        return None
    if digits.startswith("+"):
        digits = digits[1:]
    if digits.startswith("00"):
        digits = digits[2:]
    if digits.startswith("0"):  # local 07XX... -> 2567XX...
        digits = default_cc + digits[1:]
    elif len(digits) == 9 and digits[0] == "7":  # bare 7XXXXXXXX
        digits = default_cc + digits
    if not (10 <= len(digits) <= 15) or not digits.isdigit():
        return None
    return "+" + digits


def is_valid_email(raw: Any) -> str | None:
    if raw is None:
        return None
    val = str(raw).strip().lower()
    return val if _EMAIL_RE.match(val) else None


# --- Column detection (Section 3.1.1) --------------------------------------
def _detect_column(df: pd.DataFrame, hints: set[str], validator) -> str | None:
    # 1. Name matching
    for col in df.columns:
        if str(col).strip().lower() in hints:
            return col
    # 2. Value sampling — first 20 non-empty, >50% pass
    for col in df.columns:
        sample = [v for v in df[col].dropna().tolist()[:20] if str(v).strip()]
        if sample and sum(1 for v in sample if validator(v)) / len(sample) > 0.5:
            return col
    return None


def _result_from_dataframe(df: pd.DataFrame) -> ParseResult:
    df = df.dropna(how="all").reset_index(drop=True)
    df.columns = [str(c).strip() for c in df.columns]
    # Cap the rows we materialise per upload. The route already enforces a byte
    # ceiling, but a small compressed .xlsx/.docx can expand to millions of rows;
    # truncate so parsing can't exhaust worker memory (zip-bomb defence).
    total_rows = len(df)
    truncated = total_rows > _MAX_ROWS
    if truncated:
        df = df.head(_MAX_ROWS)
    result = ParseResult(columns=list(df.columns), total_rows=total_rows)
    if truncated:
        result.errors.append(
            f"Only the first {_MAX_ROWS:,} rows were processed ({total_rows:,} submitted)."
        )

    result.phone_column = _detect_column(df, _PHONE_HINTS, lambda v: normalise_phone(v) is not None)
    result.email_column = _detect_column(df, _EMAIL_HINTS, lambda v: is_valid_email(v) is not None)
    # 3. Remaining columns become merge tags
    result.merge_columns = [
        c for c in df.columns if c not in (result.phone_column, result.email_column)
    ]

    seen: set[str] = set()
    for idx, row in df.iterrows():
        raw = {c: (None if pd.isna(row[c]) else row[c]) for c in df.columns}
        phone = normalise_phone(raw.get(result.phone_column)) if result.phone_column else None
        email = is_valid_email(raw.get(result.email_column)) if result.email_column else None
        if not phone and not email:
            if len(result.errors) < _MAX_ERRORS:
                result.errors.append(f"Row {idx + 1}: no valid phone or email.")
            result.error_count += 1
            continue
        dedup_key = phone or email or ""
        if dedup_key in seen:
            result.duplicate_count += 1
            continue
        seen.add(dedup_key)
        result.contacts.append({
            "phone": phone,
            "email": email,
            "raw_data": {k: ("" if v is None else str(v)) for k, v in raw.items()},
        })
        result.valid_contacts += 1
    return result


# --- Format parsers --------------------------------------------------------
def parse_csv(content: bytes) -> ParseResult:
    encoding = chardet.detect(content).get("encoding") or "utf-8"
    text = content.decode(encoding, errors="replace")
    try:
        dialect = csv.Sniffer().sniff(text[:4096], delimiters=",;\t|")
        sep = dialect.delimiter
    except csv.Error:
        sep = ","
    df = pd.read_csv(io.StringIO(text), sep=sep, dtype=str, keep_default_na=False, na_values=[""])
    return _result_from_dataframe(df)


def parse_excel(content: bytes) -> ParseResult:
    df = pd.read_excel(io.BytesIO(content), dtype=str, sheet_name=0)  # first sheet
    return _result_from_dataframe(df)


def parse_docx(content: bytes) -> ParseResult:
    import docx  # lazy

    document = docx.Document(io.BytesIO(content))
    if document.tables:  # first table
        table = document.tables[0]
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        if rows:
            df = pd.DataFrame(rows[1:], columns=rows[0])
            return _result_from_dataframe(df)
    # Fallback: body text treated as CSV
    body = "\n".join(p.text for p in document.paragraphs if p.text.strip())
    return parse_csv(body.encode("utf-8"))


def parse_pdf(content: bytes) -> ParseResult:
    import pdfplumber  # lazy

    frames: list[pd.DataFrame] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            for table in page.extract_tables() or []:
                if table and len(table) > 1:
                    frames.append(pd.DataFrame(table[1:], columns=table[0]))
    if frames:
        return _result_from_dataframe(pd.concat(frames, ignore_index=True))
    # Camelot fallback (Section 3.1) — requires ghostscript
    try:
        import camelot  # lazy

        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".pdf") as tmp:
            tmp.write(content)
            tmp.flush()
            tables = camelot.read_pdf(tmp.name, pages="all")
        if tables.n:
            combined = pd.concat([t.df for t in tables], ignore_index=True)
            combined.columns = combined.iloc[0]
            return _result_from_dataframe(combined[1:])
    except Exception:  # noqa: BLE001
        pass
    return ParseResult(errors=["No tables found in PDF."], error_count=1)


def parse_pasted(text: str) -> ParseResult:
    """Comma/newline split. No headers — personalisation unavailable (Section 3.1)."""
    result = ParseResult(columns=[])
    seen: set[str] = set()
    tokens: list[str] = []
    for line in text.replace(",", "\n").splitlines():
        tok = line.strip()
        if tok:
            tokens.append(tok)
    result.total_rows = len(tokens)
    for i, tok in enumerate(tokens):
        phone = normalise_phone(tok)
        email = is_valid_email(tok)
        if not phone and not email:
            if len(result.errors) < _MAX_ERRORS:
                result.errors.append(f"Entry {i + 1}: '{tok}' is not a valid phone or email.")
            result.error_count += 1
            continue
        key = phone or email or ""
        if key in seen:
            result.duplicate_count += 1
            continue
        seen.add(key)
        result.contacts.append({"phone": phone, "email": email, "raw_data": {}})
        result.valid_contacts += 1
    return result


def parse_upload(content: bytes, source_format: str) -> ParseResult:
    fmt = source_format.lower().lstrip(".")
    if fmt == "csv":
        return parse_csv(content)
    if fmt in ("xlsx", "xls"):
        return parse_excel(content)
    if fmt == "docx":
        return parse_docx(content)
    if fmt == "pdf":
        return parse_pdf(content)
    raise ValueError(f"Unsupported format: {source_format}")
